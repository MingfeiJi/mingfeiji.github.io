#!/usr/bin/env python3
"""Build the site-redesign PRD DOCX from its Markdown source."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "PRD_个站全量翻新_v0.1.md"
OUTPUT = ROOT / "docs" / "PRD_个站全量翻新_v0.1.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
MUTED = "5F6B7A"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
MID_GRAY = "D8DEE8"
WHITE = "FFFFFF"
BLACK = "20242A"
GOLD = "9A6B16"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CN_FONT = "Heiti SC"


def set_run_font(run, name=CN_FONT, size=11, color=BLACK, bold=None, italic=None):
    # Leave rFonts unset so Word/LibreOffice can choose an installed CJK face.
    # Explicit Western faces in OOXML can suppress Chinese fallback in headless LO.
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def allocate_widths(rows):
    cols = len(rows[0])
    max_lengths = []
    for col in range(cols):
        length = max(len(re.sub(r"[`*_]", "", row[col])) for row in rows if col < len(row))
        max_lengths.append(max(6, min(length, 44)))
    minimum = 900 if cols <= 5 else 700
    residual = TABLE_WIDTH_DXA - minimum * cols
    weight_total = sum(max_lengths)
    widths = [minimum + int(residual * w / weight_total) for w in max_lengths]
    widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    return widths


def paragraph_border_left(paragraph, color=BLUE, size=18, space=10):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), color)
    p_bdr.append(left)


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`)")


def add_inline(paragraph, text, size=11, color=BLACK, italic=False):
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_run_font(run, size=size, color=color, italic=italic)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=color, bold=True, italic=italic)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=max(8.5, size - 0.5), color=DARK_BLUE, italic=italic)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=size, color=color, italic=italic)


def style_paragraph(paragraph, after=6, line_spacing=1.25, before=0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing
    fmt.widow_control = True


def add_body(doc, text, *, bold=False, italic=False, color=BLACK):
    p = doc.add_paragraph()
    style_paragraph(p)
    run = p.add_run(text) if bold else None
    if run is not None:
        set_run_font(run, size=11, color=color, bold=True, italic=italic)
    else:
        add_inline(p, text, size=11, color=color, italic=italic)
    return p


def add_quote(doc, text):
    p = doc.add_paragraph()
    style_paragraph(p, after=10, before=4)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.15)
    paragraph_border_left(p)
    add_inline(p, text, size=10.5, color=MUTED, italic=True)
    return p


def add_code_block(doc, lines):
    for idx, line in enumerate(lines or [""]):
        p = doc.add_paragraph()
        style_paragraph(p, after=0 if idx < len(lines) - 1 else 8, line_spacing=1.0)
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.right_indent = Inches(0.12)
        paragraph_border_left(p, color=DARK_BLUE, size=14, space=8)
        p_pr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), LIGHT_GRAY)
        p_pr.append(shd)
        run = p.add_run(line or " ")
        set_run_font(run, name="Consolas", size=8.5, color=NAVY)


def create_numbering_instance(doc):
    """Create an independent real Word numbering sequence starting at 1."""
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering
        if node.tag == qn("w:abstractNum") and node.get(qn("w:abstractNumId"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering
        if node.tag == qn("w:num") and node.get(qn("w:numId"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    level.append(level_text)
    level_jc = OxmlElement("w:lvlJc")
    level_jc.set(qn("w:val"), "left")
    level.append(level_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.append(indent)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_list_item(doc, text, numbered=False, level=0, num_id=None):
    style_name = "List Paragraph" if numbered else "List Bullet"
    p = doc.add_paragraph(style=style_name)
    style_paragraph(p, after=4)
    p.paragraph_format.left_indent = Inches(0.375 + level * 0.22)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    if numbered:
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), str(level))
        num_pr.append(ilvl)
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), str(num_id))
        num_pr.append(num_id_node)
        p._p.get_or_add_pPr().append(num_pr)
    add_inline(p, text, size=10.7)
    return p


def add_compact_records(doc, header, data):
    for row_index, row in enumerate(data, start=1):
        p = doc.add_paragraph()
        style_paragraph(p, before=8, after=3)
        run = p.add_run(f"记录 {row_index}")
        set_run_font(run, size=10.5, color=DARK_BLUE, bold=True)
        for key, value in zip(header, row):
            item = doc.add_paragraph(style="List Bullet")
            style_paragraph(item, after=2)
            item.paragraph_format.left_indent = Inches(0.38)
            item.paragraph_format.first_line_indent = Inches(-0.188)
            label = item.add_run(f"{key}：")
            set_run_font(label, size=9.8, color=MUTED, bold=True)
            add_inline(item, value, size=9.8)


def add_table(doc, rows):
    header, data = rows[0], rows[1:]
    if len(header) > 6:
        add_compact_records(doc, header, data)
        return

    widths = allocate_widths(rows)
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    set_row_cant_split(header_row)

    for idx, text in enumerate(header):
        cell = header_row.cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        style_paragraph(p, after=0, line_spacing=1.05)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline(p, text, size=9.2, color=NAVY)
        for run in p.runs:
            run.bold = True

    for row_index, row in enumerate(data):
        cells = table.add_row().cells
        set_row_cant_split(table.rows[-1])
        for idx, text in enumerate(row):
            cell = cells[idx]
            if row_index % 2 == 1:
                set_cell_shading(cell, "FAFBFC")
            p = cell.paragraphs[0]
            style_paragraph(p, after=0, line_spacing=1.08)
            if len(header) <= 3 and idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif len(text) <= 12:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline(p, text, size=8.9 if len(header) >= 5 else 9.3)

    set_table_geometry(table, widths)
    spacer = doc.add_paragraph()
    style_paragraph(spacer, after=5)


def set_document_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    token_map = {
        "Title": (29, NAVY, 0, 8),
        "Subtitle": (13, MUTED, 0, 18),
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in token_map.items():
        style = doc.styles[style_name]
        style.font.size = Pt(size)
        style.font.bold = style_name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True


def add_field(paragraph, field_code):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_code
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_sep, text, fld_char_end])
    set_run_font(run, size=8.5, color=MUTED)


def add_page_furniture(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_paragraph(header, after=0, line_spacing=1.0)
    left = header.add_run("SITE REDESIGN PRD")
    set_run_font(left, size=8.5, color=MUTED, bold=True)
    right = header.add_run("    ·    MINGFEI JI")
    set_run_font(right, size=8.5, color=GOLD)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_paragraph(footer, after=0, line_spacing=1.0)
    label = footer.add_run("PRD v0.1  ·  ")
    set_run_font(label, size=8.5, color=MUTED)
    add_field(footer, "PAGE")


def add_cover(doc):
    p = doc.add_paragraph()
    style_paragraph(p, before=72, after=14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PRODUCT / DESIGN / DELIVERY")
    set_run_font(run, size=10, color=GOLD, bold=True)

    title = doc.add_paragraph()
    style_paragraph(title, after=8, line_spacing=1.0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("纪鸣飞个人站全量翻新")
    set_run_font(title_run, size=29, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    style_paragraph(subtitle, after=18, line_spacing=1.0)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("三主题设计系统 · Vibe Coding 专区 · 人物照片 · Vercel / Supabase 路线")
    set_run_font(subtitle_run, size=13, color=MUTED, italic=True)

    p = doc.add_paragraph()
    style_paragraph(p, before=18, after=8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PRD v0.1  /  2026-08-01")
    set_run_font(run, size=10.5, color=MUTED, bold=True)

    lead = doc.add_paragraph()
    style_paragraph(lead, before=34, after=26, line_spacing=1.35)
    lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead.paragraph_format.left_indent = Inches(0.6)
    lead.paragraph_format.right_indent = Inches(0.6)
    run = lead.add_run("从“内容丰富的长简历”升级为“人的构建现场”：让访客先认识纪鸣飞，再验证作品、方法和 Vibe Coding 系统。")
    set_run_font(run, size=13, color=NAVY, bold=True)

    meta = doc.add_table(rows=5, cols=2)
    meta.style = "Table Grid"
    rows = [
        ("元信息", "内容"),
        ("项目", "MingfeiJi/mingfeiji.github.io"),
        ("工作区", "/Users/temptrip/sites/mingfeiji-site-redesign"),
        ("分支", "agent/site-redesign"),
        ("状态", "规划初稿 / 等待方向与素材确认"),
    ]
    widths = [1900, 7460]
    set_repeat_table_header(meta.rows[0])
    for row_index, (row, (label, value)) in enumerate(zip(meta.rows, rows)):
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        if row_index == 0:
            set_cell_shading(row.cells[1], LIGHT_BLUE)
        for idx, text in enumerate((label, value)):
            p = row.cells[idx].paragraphs[0]
            style_paragraph(p, after=0, line_spacing=1.05)
            add_inline(p, text, size=9.8, color=NAVY if idx == 0 else BLACK)
            if idx == 0 or row_index == 0:
                p.runs[0].bold = True
    set_table_geometry(meta, widths)
    doc.add_page_break()


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
        rows.append(cells)
        i += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", c) for c in rows[1]):
        rows.pop(1)
    return rows, i


def build():
    source = SOURCE.read_text(encoding="utf-8")
    lines = source.splitlines()

    doc = Document()
    set_document_styles(doc)
    add_page_furniture(doc)
    add_cover(doc)

    in_code = False
    code_lines = []
    skip_first_title = True
    current_num_id = None
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()

        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(raw.rstrip())
            i += 1
            continue

        if not line or line == "---":
            current_num_id = None
            i += 1
            continue

        if line.startswith("# ") and skip_first_title:
            current_num_id = None
            skip_first_title = False
            i += 1
            continue

        if line.startswith("## "):
            current_num_id = None
            heading = doc.add_paragraph(style="Heading 1")
            add_inline(heading, line[3:], size=16, color=BLUE)
            for run in heading.runs:
                run.bold = True
            i += 1
            continue
        if line.startswith("### "):
            current_num_id = None
            heading = doc.add_paragraph(style="Heading 2")
            add_inline(heading, line[4:], size=13, color=BLUE)
            for run in heading.runs:
                run.bold = True
            i += 1
            continue
        if line.startswith("#### "):
            current_num_id = None
            heading = doc.add_paragraph(style="Heading 3")
            add_inline(heading, line[5:], size=12, color=DARK_BLUE)
            for run in heading.runs:
                run.bold = True
            i += 1
            continue
        if line.startswith("# "):
            current_num_id = None
            heading = doc.add_paragraph(style="Heading 1")
            add_inline(heading, line[2:], size=16, color=BLUE)
            i += 1
            continue

        if line.startswith("|"):
            current_num_id = None
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue

        if line.startswith("> "):
            current_num_id = None
            add_quote(doc, line[2:])
            i += 1
            continue

        number_match = re.match(r"^(\d+)\.\s+(.*)$", line)
        if number_match:
            if current_num_id is None:
                current_num_id = create_numbering_instance(doc)
            add_list_item(doc, number_match.group(2), numbered=True, num_id=current_num_id)
            i += 1
            continue

        if re.match(r"^- \[[ xX]\]\s+", line):
            current_num_id = None
            text = re.sub(r"^- \[[ xX]\]\s+", "", line)
            add_list_item(doc, f"☐ {text}")
            i += 1
            continue

        if line.startswith("- "):
            current_num_id = None
            add_list_item(doc, line[2:])
            i += 1
            continue

        current_num_id = None
        add_body(doc, line)
        i += 1

    core = doc.core_properties
    core.title = "纪鸣飞个人站全量翻新 PRD v0.1"
    core.subject = "三主题设计系统、Vibe Coding 专区、个人照片与部署规划"
    core.author = "Codex"
    core.keywords = "Mingfei Ji, personal site, Vibe Coding, Astro, Vercel, Supabase"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    try:
        build()
    except Exception as exc:
        print(f"build failed: {exc}", file=sys.stderr)
        raise
